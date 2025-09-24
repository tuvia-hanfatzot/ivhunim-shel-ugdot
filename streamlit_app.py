import io
import re
import zipfile
from collections import defaultdict
from typing import Iterable, List, Optional, Tuple
import pandas as pd
import streamlit as st

# --- ID detection: 9 digits starting with 4, 8, or 9 ---
ID_PATTERN = re.compile(r'\b([489]\d{8})\b')


def extract_ids_from_text(text: str) -> List[str]:
    return list(dict.fromkeys(ID_PATTERN.findall(text or "")))  # preserve order & dedupe


# --- File readers ---
def read_docx_bytes(data: bytes) -> str:
    """Extract text from a .docx file (paragraphs + tables)."""
    try:
        from docx import Document  # python-docx
    except Exception as e:
        raise RuntimeError("python-docx is required to parse .docx files") from e

    file_like = io.BytesIO(data)
    doc = Document(file_like)

    parts = []
    for p in doc.paragraphs:
        parts.append(p.text)
    # Extract table text as well
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                parts.append(cell.text)

    return "\n".join(parts)


def read_xlsx_bytes(data: bytes) -> str:
    """Extract all cell values from an .xlsx workbook using openpyxl."""
    try:
        from openpyxl import load_workbook
    except Exception as e:
        raise RuntimeError("openpyxl is required to parse .xlsx files") from e

    wb = load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    parts = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            for val in row:
                if val is None:
                    continue
                parts.append(str(val))
    return "\n".join(parts)


def read_xls_bytes(data: bytes) -> str:
    """Extract all cell values from an .xls workbook using xlrd (v2 supports .xls only)."""
    try:
        import xlrd  # Ensure version supports .xls
    except Exception as e:
        raise RuntimeError("xlrd is required (and must support .xls) to parse .xls files") from e

    book = xlrd.open_workbook(file_contents=data)
    parts = []
    for sheet in book.sheets():
        for r in range(sheet.nrows):
            for c in range(sheet.ncols):
                val = sheet.cell_value(r, c)
                if val in (None, ""):
                    continue
                parts.append(str(val))
    return "\n".join(parts)

def read_doc_bytes(data: bytes) -> str:
    """
    Best-effort text extraction for legacy .doc.
    Tries 'textract' if available; otherwise raises with a clear message.
    """
    try:
        import textract  # requires system dependencies
    except Exception as e:
        raise RuntimeError(
            "textract is required to parse legacy .doc files and may need system packages (antiword/catdoc)."
        ) from e

    txt = textract.process(io.BytesIO(data), extension="doc")  # type: ignore
    try:
        return txt.decode("utf-8", errors="replace")
    except Exception:
        return str(txt)

# --- Helpers for ZIP traversal ---
def nearest_folder_from_path(path: str) -> Optional[str]:
    """
    Return the nearest enclosing folder (last folder before the filename),
    or None if the file is at the ZIP root.
    """
    segments = [seg for seg in path.split("/") if seg and seg != "."]
    return segments[-2] if len(segments) >= 2 else None

def extract_folder_file_pairs(zf: zipfile.ZipFile) -> Iterable[Tuple[str, zipfile.ZipInfo]]:
    """Yield (folder_name, zipinfo) for supported files under their nearest enclosing folder (or ROOT)."""
    for info in zf.infolist():
        if info.is_dir():
            continue
        path = info.filename.replace("\\", "/")
        folder = nearest_folder_from_path(path) or "ROOT"
        ext = path.lower().rsplit(".", 1)[-1] if "." in path else ""
        if ext in {"docx", "doc", "xls", "xlsx"}:
            yield folder, info

def harvest_ids_from_zip(zf: zipfile.ZipFile) -> Tuple[pd.DataFrame, List[str]]:
    """
    Return (DataFrame with columns: Folder, ID, file_order, row_order, warnings).
    - Preserves file traversal order within each folder (file_order).
    - Preserves ID appearance order within each file (row_order).
    - De-duplicates per (Folder, ID) keeping the earliest by file_order then row_order.
    """
    rows = []
    warnings: List[str] = []

    # Track per-folder file order in the sequence we encounter files in the ZIP
    folder_file_order: dict[str, dict[str, int]] = {}

    for folder, info in extract_folder_file_pairs(zf):
        path = info.filename.replace("\\", "/")
        filename = path.split("/")[-1]
        ext = path.lower().rsplit(".", 1)[-1]

        try:
            data = zf.read(info)
        except Exception as e:
            warnings.append(f"Could not read '{path}': {e}")
            continue

        text = ""
        try:
            if ext == "docx":
                text = read_docx_bytes(data)
            elif ext == "doc":
                text = read_doc_bytes(data)  # may warn if textract missing
            elif ext == "xlsx":
                text = read_xlsx_bytes(data)
            elif ext == "xls":
                text = read_xls_bytes(data)
        except Exception as e:
            warnings.append(f"Could not parse '{path}': {e}")
            continue

        ids = extract_ids_from_text(text)
        if not ids:
            warnings.append(f"No IDs found in '{path}'.")
            continue

        # Assign/order this file within its folder
        folder_file_order.setdefault(folder, {})
        if filename not in folder_file_order[folder]:
            folder_file_order[folder][filename] = len(folder_file_order[folder])  # 0,1,2,...

        file_order = folder_file_order[folder][filename]

        # Record each ID with “row_order” = sequence within this file
        for row_order, pid in enumerate(ids):
            rows.append({
                "Folder": folder,
                "ID": pid,
                "file_order": file_order,
                "row_order": row_order,
                # keep filename only internally for tie-breaks if ever needed
                "_source": filename,
            })

    if not rows:
        df = pd.DataFrame(columns=["Folder", "ID", "file_order", "row_order"])
        return df, warnings

    df = pd.DataFrame(rows)

    # De-dupe per (Folder, ID) keeping earliest by file_order -> row_order -> first seen
    df = (
        df.sort_values(["Folder", "file_order", "row_order", "_source"])
          .drop_duplicates(subset=["Folder", "ID"], keep="first")
          .reset_index(drop=True)
    )

    # No need to expose source filename
    df = df.drop(columns=["_source"])
    return df, warnings

from openpyxl import load_workbook, Workbook
from openpyxl.styles import PatternFill, Font
from openpyxl.utils import get_column_letter
from openpyxl.styles import PatternFill, Font, Alignment

# Style constants
HEADER_ORANGE = PatternFill(start_color="FFC000", end_color="FFC000", fill_type="solid")  # Excel orange
BLACK_FILL    = PatternFill(start_color="000000", end_color="000000", fill_type="solid")
CENTER_ALIGN  = Alignment(horizontal="center", vertical="center", wrap_text=False)

# Column layout (RTL): A..F
COL_MSD          = "A"  # מס"ד (numeric/string, replaces Folder)
COL_SHEM_TIAUM   = "B"  # שם תיאום (left empty per request)
COL_TZ_MISHTATF  = "C"  # ת"ז משתתפים (replaces ID)
COL_PNIYA        = "D"  # פנייה (empty)
COL_TAFKID       = "E"  # תפקיד (רק למי שמסורב) (empty)
COL_MODIIN       = "F"  # תגובת מודיעין (empty)

HEADERS_RTL = ["מס\"ד", "שם תיאום", "ת\"ז משתתפים", "פנייה", "תפקיד (רק למי שמסורב)", "תגובת מודיעין"]

YELLOW_FILL = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
BLACK_FILL  = PatternFill(start_color="000000", end_color="000000", fill_type="solid")

def _sanitize_sheet_title(name: str) -> str:
    # Excel sheet title constraints
    bad = set('[]:*?/\\')
    clean = "".join(ch for ch in str(name) if ch not in bad).strip()
    return (clean or "Folder")[:31]

def _is_solid_fill(cell) -> bool:
    try:
        return cell.fill is not None and cell.fill.fill_type == "solid" and (
            (cell.fill.start_color.rgb or "").upper().endswith("000000")  # black
        )
    except Exception:
        return False

def _is_separator_row(ws, row_idx: int) -> bool:
    # a separator = full black row across A:F
    cols = ["A", "B", "C", "D", "E", "F"]
    return all(_is_solid_fill(ws[f"{c}{row_idx}"]) for c in cols)

def _parse_summary(ws_summary):
    """
    Parse the 'Summary' sheet in RTL 6-col layout:
    A: מס"ד      (folder)
    B: שם תיאום  (blank)
    C: ת"ז משתתפים (ID)
    D: פנייה      (blank)
    E: תפקיד...   (blank)
    F: תגובת מודיעין (blank)

    Returns blocks: [{"folder": str, "start": int, "end": int, "sep": int|None, "ids": [str, ...]}]
    """
    blocks = []
    max_row = ws_summary.max_row
    r = 2  # skip header
    while r <= max_row:
        # skip any leading separators/empty
        while r <= max_row and _is_separator_row(ws_summary, r):
            r += 1
        while r <= max_row and all(ws_summary[f"{c}{r}"].value in (None, "") for c in ["A","B","C","D","E","F"]):
            r += 1
        if r > max_row:
            break

        # start of block
        folder = ws_summary[f"A{r}"].value
        if not folder and r > 2:
            rr = r - 1
            while rr >= 2 and ws_summary[f"A{rr}"].value in (None, ""):
                rr -= 1
            folder = ws_summary[f"A{rr}"].value

        start = r
        ids = []
        while r <= max_row and not _is_separator_row(ws_summary, r):
            id_val = ws_summary[f"C{r}"].value  # IDs live in col C
            if id_val not in (None, ""):
                ids.append(str(id_val))
            r += 1
        end = r - 1
        sep = r if (r <= max_row and _is_separator_row(ws_summary, r)) else None
        if sep:
            r += 1
        blocks.append({"folder": str(folder) if folder is not None else "", "start": start, "end": end, "sep": sep, "ids": ids})
    return blocks

def _ensure_provenance_sheet(wb):
    # no-op here; kept for symmetry if you add provenance later
    return

def _append_ids_to_folder_sheet(wb, folder: str, new_ids: list[str]):
    title = _sanitize_sheet_title(folder)
    if title in wb.sheetnames:
        ws = wb[title]
        # append below last row
        for pid in new_ids:
            ws.append([pid])
    else:
        ws = wb.create_sheet(title=title)
        ws.append(["ID"])
        for pid in new_ids:
            ws.append([pid])
        ws.column_dimensions["A"].width = 20
        ws.freeze_panes = "A2"

def update_existing_workbook(existing_xlsx: bytes, update_df: pd.DataFrame) -> bytes:
    """
    Update the RTL Summary sheet (A..F) where only A=מס\"ד and C=ת\"ז משתתפים are editable.
    New IDs for an existing מס\"ד are inserted just before its separator row.
    New מס\"ד blocks are appended at the bottom.
    Newly added rows are highlighted in yellow (A..F for those rows).
    """
    wb = load_workbook(io.BytesIO(existing_xlsx))
    if "Summary" not in wb.sheetnames:
        raise RuntimeError("The uploaded Excel has no 'Summary' sheet.")

    ws = wb["Summary"]
    ws.sheet_view.rightToLeft = True  # keep RTL if missing
    blocks = _parse_summary(ws)

    folder_to_ids = {b["folder"]: set(b["ids"]) for b in blocks}
    folder_order  = [b["folder"] for b in blocks]

    # Prepare additions from update_df
    upd_sorted = update_df.sort_values(["Folder", "file_order", "row_order"])
    additions: dict[str, list[str]] = {}
    for folder, sub in upd_sorted.groupby("Folder", sort=False):
        have = folder_to_ids.get(folder, set())
        to_add = [str(pid) for pid in sub["ID"].astype(str).tolist() if pid not in have]
        if to_add:
            additions[folder] = to_add

    if not additions:
        out = io.BytesIO()
        wb.save(out)
        return out.getvalue()

    # Bottom-up insert for existing folders
    blocks_by_folder = {b["folder"]: b for b in blocks}
    for folder in reversed(folder_order):
        if folder not in additions:
            continue
        new_ids = additions[folder]
        b = blocks_by_folder[folder]
        insert_at = (b["sep"] if b["sep"] else b["end"] + 1)

        ws.insert_rows(insert_at, amount=len(new_ids))
        r = insert_at
        for idx, pid in enumerate(new_ids):
            # Populate only A (מס"ד) for the first row (to keep compact display),
            # and C (ת"ז משתתפים) for each row. Other columns remain blank.
            ws[f"{COL_MSD}{r}"].value = folder if idx == 0 else ""
            ws[f"{COL_SHEM_TIAUM}{r}"].value = ""
            ws[f"{COL_TZ_MISHTATF}{r}"].value = pid
            ws[f"{COL_PNIYA}{r}"].value = ""
            ws[f"{COL_TAFKID}{r}"].value = ""
            ws[f"{COL_MODIIN}{r}"].value = ""
            # Yellow highlight & center alignment across A:F
            for col in ["A","B","C","D","E","F"]:
                cell = ws[f"{col}{r}"]
                cell.fill = YELLOW_FILL
                cell.alignment = CENTER_ALIGN
            r += 1

        # shift indices for blocks following
        shift = len(new_ids)
        for bb in blocks:
            if bb["start"] >= insert_at:
                bb["start"] += shift
            if bb["end"] >= insert_at:
                bb["end"] += shift
            if bb["sep"] and bb["sep"] >= insert_at:
                bb["sep"] += shift

        folder_to_ids[folder].update(new_ids)

    # New folders: append at bottom
    new_folders = [f for f in additions.keys() if f not in blocks_by_folder]
    if new_folders:
        last_row = ws.max_row
        for folder in new_folders:
            new_ids = additions[folder]
            if not new_ids:
                continue
            row = last_row + 1
            # first row with מס"ד + first ID
            ws[f"{COL_MSD}{row}"].value = folder
            ws[f"{COL_SHEM_TIAUM}{row}"].value = ""
            ws[f"{COL_TZ_MISHTATF}{row}"].value = new_ids[0]
            ws[f"{COL_PNIYA}{row}"].value = ""
            ws[f"{COL_TAFKID}{row}"].value = ""
            ws[f"{COL_MODIIN}{row}"].value = ""
            for col in ["A","B","C","D","E","F"]:
                cell = ws[f"{col}{row}"]
                cell.fill = YELLOW_FILL
                cell.alignment = CENTER_ALIGN
            row += 1
            # remaining IDs
            for pid in new_ids[1:]:
                for col in ["A","B","C","D","E","F"]:
                    ws[f"{col}{row}"].value = ""
                ws[f"{COL_TZ_MISHTATF}{row}"].value = pid
                for col in ["A","B","C","D","E","F"]:
                    cell = ws[f"{col}{row}"]
                    cell.fill = YELLOW_FILL
                    cell.alignment = CENTER_ALIGN
                row += 1
            # separator row
            for col in ["A","B","C","D","E","F"]:
                ws[f"{col}{row}"].value = ""
                ws[f"{col}{row}"].fill = BLACK_FILL
                ws[f"{col}{row}"].alignment = CENTER_ALIGN
            last_row = row

    # Keep header orange + centered (in case the source wasn't)
    for i in range(1, 7):
        h = ws.cell(row=1, column=i)
        h.fill = HEADER_ORANGE
        h.font = Font(bold=True)
        h.alignment = CENTER_ALIGN

    # Ensure RTL and centre alignment across body
    ws.sheet_view.rightToLeft = True
    for r in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=6):
        for c in r:
            if c.fill is None or c.row == 1:
                c.alignment = CENTER_ALIGN

    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()

def to_excel_bytes(df: pd.DataFrame) -> bytes:
    """
    Create Excel with RTL 'Summary' sheet and columns:
    A מס"ד | B שם תיאום | C ת"ז משתתפים | D פנייה | E תפקיד (רק למי שמסורב) | F תגובת מודיעין
    Only A (מס"ד) and C (ת"ז משתתפים) are filled. Others remain empty.
    All cells centered; header row orange.
    A black full-width separator row after each מס"ד block.
    """
    from openpyxl import Workbook

    df_sorted = df.sort_values(["Folder", "file_order", "row_order"]).reset_index(drop=True)

    wb = Workbook()
    ws = wb.active
    ws.title = "Summary"
    ws.sheet_view.rightToLeft = True  # RTL sheet

    # Header
    ws.append(HEADERS_RTL)
    for i, _ in enumerate(HEADERS_RTL, start=1):
        cell = ws.cell(row=1, column=i)
        cell.fill = HEADER_ORANGE
        cell.font = Font(bold=True)
        cell.alignment = CENTER_ALIGN

    # Body
    for folder, sub in df_sorted.groupby("Folder", sort=True):
        sub = sub.sort_values(["file_order", "row_order"])
        first = True
        for _, row in sub.iterrows():
            values = ["", "", "", "", "", ""]
            # Only A (מס"ד) and C (ת"ז משתתפים) are populated
            values[0] = folder if first else ""         # A
            values[2] = str(row["ID"])                  # C
            ws.append(values)
            # center alignment for the whole new row
            r = ws.max_row
            for col_idx in range(1, 7):
                ws.cell(row=r, column=col_idx).alignment = CENTER_ALIGN
            first = False

        # separator row across A:F in black
        sep_row_idx = ws.max_row + 1
        ws.append(["", "", "", "", "", ""])
        for col in range(1, 7):
            ws.cell(row=sep_row_idx, column=col).fill = BLACK_FILL

    # Column widths (simple auto)
    for col_cells in ws.columns:
        length = max(len(str(c.value)) if c.value is not None else 0 for c in col_cells)
        ws.column_dimensions[col_cells[0].column_letter].width = min(max(length + 2, 12), 40)

    ws.freeze_panes = "A2"

    # Ensure all body cells are centered (in case)
    for r in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=6):
        for c in r:
            c.alignment = CENTER_ALIGN

    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()

# --- Streamlit UI ---
st.set_page_config(page_title="ID Harvester to Excel", page_icon="📄", layout="wide")
st.title("📄 ID Harvester → Excel")
st.caption(
    "Upload a ZIP of your folder. Each numbered subfolder may contain .docx, .doc, .xls, or .xlsx files. "
    "The app extracts 9-digit IDs beginning with 4/8/9 and groups them by folder."
)

st.subheader("Mode")
tab_create, tab_update = st.tabs(["Create from ZIP", "Update existing Excel with ZIP"])

with tab_create:
    base_zip = st.file_uploader("Upload a ZIP archive (create new Excel)", type=["zip"], key="create_zip")
    if base_zip is not None:
        try:
            with zipfile.ZipFile(base_zip) as zf:
                df, warnings = harvest_ids_from_zip(zf)

            if df.empty:
                st.warning("No IDs were found. Please check your ZIP structure and file contents.")
            else:
                st.subheader("Preview")
                st.dataframe(df, use_container_width=True)

                excel_bytes = to_excel_bytes(df)
                st.download_button(
                    label="⬇️ Download Excel",
                    data=excel_bytes,
                    file_name="ids_by_folder.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )

            if warnings:
                st.divider()
                st.subheader("Warnings")
                for w in warnings:
                    st.write("• " + w)

        except zipfile.BadZipFile:
            st.error("The uploaded file is not a valid ZIP archive.")
        except Exception as e:
            st.exception(e)

with tab_update:
    existing_xlsx = st.file_uploader("Upload the existing Excel to update (.xlsx)", type=["xlsx"], key="existing_xlsx")
    update_zip    = st.file_uploader("Upload the update ZIP archive", type=["zip"], key="update_zip")

    if existing_xlsx is not None and update_zip is not None:
        try:
            with zipfile.ZipFile(update_zip) as zf:
                upd_df, upd_warnings = harvest_ids_from_zip(zf)

            if upd_df.empty:
                st.warning("No IDs were found in the update ZIP.")
            else:
                # Show what will be applied (per folder unique new IDs after comparing will be computed inside updater)
                st.subheader("Update ZIP Preview")
                st.dataframe(upd_df, use_container_width=True)

                updated_bytes = update_existing_workbook(existing_xlsx.read(), upd_df)
                st.download_button(
                    label="⬇️ Download Updated Excel",
                    data=updated_bytes,
                    file_name="ids_by_folder_updated.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )

            if upd_warnings:
                st.divider()
                st.subheader("Warnings (update ZIP)")
                for w in upd_warnings:
                    st.write("• " + w)

        except zipfile.BadZipFile:
            st.error("The update file is not a valid ZIP archive.")
        except Exception as e:
            st.exception(e)
